from model.user_model import UserModel
from tkinter import messagebox
import docx
from datetime import datetime

class PrintListController:
    def __init__(self, view):
        self.view = view


    def search_user(self, ma_gia_dinh):
        user_model = UserModel()
        return user_model.search(ma_gia_dinh)
    
    def list_printed(self, list_ma):
        user_model = UserModel()
        return user_model.search_list(list_ma)

    def ngay_print(self, filename, doc_list):
        document = docx.Document()
        for doc in doc_list:
        # Kiểm tra nếu doc là None hoặc không phải là từ điển
            if doc is None or not isinstance(doc, dict):
                continue  # Bỏ qua phần tử này và tiếp tục với phần tử tiếp theo
            para = document.add_paragraph()
            # para.add_run("Thành Viên: ")
            para.add_run(f"TC: {doc.get('tinChu')}\n")
            for member in doc.get("thanhVien", []):
                para.add_run(f"{member.get('ho_va_ten')} ")
        document.save(filename)
        messagebox.showinfo("Thông báo", f"Xuất file 1/15 thành công ra màn hình máy tính")

    def vu_lan_print(self, filename, doc_list):
        document = docx.Document()
        for doc in doc_list:
        # Kiểm tra nếu doc là None hoặc không phải là từ điển
            if doc is None or not isinstance(doc, dict):
                continue  # Bỏ qua phần tử này và tiếp tục với phần tử tiếp theo
            para = document.add_paragraph()
            para.add_run(f"TC: {doc.get('tinChu')}\n")
            para.add_run(f"CL: {doc.get('chanLinh')}\n")
        document.save(filename)
        messagebox.showinfo("Thông báo", f"Xuất file VL thành công ra màn hình máy tính")

    def dau_nam_print(self, filename, doc_list):
        document = docx.Document()
        for doc in doc_list:
        # Kiểm tra nếu doc là None hoặc không phải là từ điển
            if doc is None or not isinstance(doc, dict):
                continue  # Bỏ qua phần tử này và tiếp tục với phần tử tiếp theo
            para = document.add_paragraph()
            para.add_run(f"TC: {doc.get('tinChu')}\n")
            for member in doc.get("thanhVien", []):
                current_year = datetime.now().year 
                nam_sinh = member.get("nam_sinh", "Không rõ").strip()
                if nam_sinh.isdigit() and len(nam_sinh) == 4:
                    age = current_year - int(nam_sinh)
                else:
                    age = "ko"
                para.add_run(f"{member.get('ho_va_ten')} ({age} tuổi), ")
        document.save(filename)
        messagebox.showinfo("Thông báo", f"Xuất file DN thành công ra màn hình máy tính")
        
    def con_ban_print(self, filename, doc_list):
        document = docx.Document()
        for doc in doc_list:
        # Kiểm tra nếu doc là None hoặc không phải là từ điển
            if doc is None or not isinstance(doc, dict):
                continue  # Bỏ qua phần tử này và tiếp tục với phần tử tiếp theo
            para = document.add_paragraph()
            para.add_run(f"TC: {doc.get('tinChu')}\n")
            # para.add_run("Thành Viên: ")
            for member in doc.get("thanhVien", []):
                if member.get('con_ban'):
                    para.add_run(f"{member.get('ho_va_ten')} ")
        document.save(filename)
        messagebox.showinfo("Thông báo", f"Xuất file CB thành công ra màn hình máy tính")

    def sao_print(self, filename, doc_list):
        document = docx.Document()
        for doc in doc_list:
        # Kiểm tra nếu doc là None hoặc không phải là từ điển
            if doc is None or not isinstance(doc, dict):
                continue  # Bỏ qua phần tử này và tiếp tục với phần tử tiếp theo
            para = document.add_paragraph()
            para.add_run(f"TC: {doc.get('tinChu')}\n")
            for member in doc.get("thanhVien", []):
                current_year = datetime.now().year 
                nam_sinh = member.get("nam_sinh", "Không rõ").strip()
                if nam_sinh.isdigit() and len(nam_sinh) == 4:
                    age = current_year - int(nam_sinh)
                else:
                    age = "-1"
                
                sao_nam = ["La Hầu", "Thổ Tú", "Thủy Diệu", "Thái Bạch", "Thái Dương", "Văn Hớn", "Kế Đô", "Thái Âm", "Mộc Đức"]
                sao_nu = ["Kế Đô", "Văn Hớn", "Mộc Đức", "Thái Âm", "Thổ Tú", "La Hầu", "Thái Dương", "Thái Bạch", "Thủy Diệu"]
                nam_nu = member.get("nam_nu").strip()
                if nam_nu == "xx":
                    sao = sao_nu[(age - 1) % 9]
                if nam_nu == "xy":
                    sao = sao_nam[(age - 1) % 9]
                para.add_run(f"{member.get('ho_va_ten')} ({age} tuổi - {sao}), ")
        document.save(filename)
        messagebox.showinfo("Thông báo", f"Xuất file SAO thành công ra màn hình máy tính")